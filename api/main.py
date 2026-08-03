"""
固件漏洞扫描平台 - FastAPI Web API
支持批量扫描和任务队列
"""

import os
import sys
from pathlib import Path
from typing import List, Optional
from datetime import datetime
import logging
import json

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Query, Request
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from scanner.engine import FirmwareExtractor, SBOMGenerator, CVEMatcher, Vulnerability
from scanner.task_queue import get_scan_queue, ScanTask, TaskStatus, ScanQueue
import yaml

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 加载配置
config_path = Path(__file__).parent.parent / "config.yaml"
with open(config_path, encoding='utf-8') as f:
    config = yaml.safe_load(f)

# 处理环境变量占位符 (如 ${GRYPE_DB_PATH:~/.local/share/grype/grype.db})
def resolve_env_var(value):
    """解析环境变量占位符"""
    if not isinstance(value, str):
        return value
    
    import re
    pattern = r'\$\{([^}:]+)(?::(.+))?\}'
    match = re.match(pattern, value)
    
    if match:
        env_name = match.group(1)
        default_value = match.group(2)
        return os.environ.get(env_name, default_value or '')
    
    return value

# 递归处理配置中的所有字符串值
def process_config_values(cfg):
    """递归处理配置中的环境变量"""
    if isinstance(cfg, dict):
        return {k: process_config_values(v) for k, v in cfg.items()}
    elif isinstance(cfg, list):
        return [process_config_values(item) for item in cfg]
    else:
        return resolve_env_var(cfg)

config = process_config_values(config)

# 展开 ~ 符号为家目录
if 'paths' in config:
    for key in config['paths']:
        path_val = config['paths'][key]
        if isinstance(path_val, str) and path_val.startswith('~'):
            config['paths'][key] = str(Path(path_val).expanduser())

logger.info(f"Grype DB 路径：{config['paths'].get('grype_db', '未配置')}")

app = FastAPI(title="固件漏洞扫描平台", version="2.1-alpha")

# 模板和静态文件
templates_path = Path(__file__).parent.parent / "frontend" / "templates"
static_path = Path(__file__).parent.parent / "frontend" / "static"

templates = Jinja2Templates(directory=str(templates_path))

# 确保目录存在
for dir_path in [config['paths']['uploads'], 
                 config['paths']['workspace'],
                 config['paths']['reports']]:
    Path(dir_path).mkdir(parents=True, exist_ok=True)

# 初始化全局扫描队列（根据配置设置并发数）
MAX_CONCURRENT = config.get('queue', {}).get('max_concurrent', 3)
scan_queue: Optional[ScanQueue] = None

def get_queue() -> ScanQueue:
    """获取扫描队列实例"""
    global scan_queue
    if scan_queue is None:
        scan_queue = get_scan_queue(max_concurrent=MAX_CONCURRENT)
    return scan_queue


# Pydantic 模型
class ScanRequest(BaseModel):
    firmware_id: str
    firmware_type: str

class BatchScanRequest(BaseModel):
    files: List[dict]  # [{"path": "...", "type": "..."}, ...]
    
class TaskStatusResponse(BaseModel):
    task_id: str
    filename: str
    status: str
    progress: int
    created_at: str
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    error_message: Optional[str] = None

class QueueStatsResponse(BaseModel):
    total: int
    pending: int
    queued: int
    running: int
    completed: int
    failed: int
    active_workers: int
    max_concurrent: int

# 全局任务队列（用于向后兼容）
scan_results_store = {}


@app.on_event("startup")
async def startup_event():
    """启动时初始化队列"""
    logger.info("启动扫描队列服务...")
    queue = get_queue()
    queue.start()
    logger.info(f"扫描队列已启动 (最大并发：{queue.max_concurrent})")


@app.on_event("shutdown")
async def shutdown_event():
    """关闭时清理队列"""
    global scan_queue
    if scan_queue:
        logger.info("关闭扫描队列...")
        scan_queue.close()
        logger.info("扫描队列已关闭")


@app.get("/")
async def root(request: Request):
    """首页"""
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/api/dashboard")
async def dashboard_page(request: Request):
    """仪表板页面（旧版本兼容）"""
    return templates.TemplateResponse("dashboard.html", {"request": request})


@app.post("/api/upload")
async def upload_firmware(file: UploadFile = File(...)):
    """上传固件文件"""
    try:
        file_path = Path(config['paths']['uploads']) / file.filename
        
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        firmware_id = Path(file.filename).stem
        scan_results_store[firmware_id] = {
            'path': str(file_path),
            'filename': file.filename,
            'upload_time': datetime.now().isoformat(),
            'status': 'uploaded'
        }
        
        return {
            "success": True,
            "firmware_id": firmware_id,
            "message": f"固件已上传：{file.filename}",
            "path": str(file_path)
        }
    except Exception as e:
        logger.error(f"上传失败：{e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/scan")
async def scan_firmware_single(firmware_id: str = Form(...), 
                               firmware_type: str = Form(...)):
    """执行单个固件扫描（向后兼容）"""
    try:
        if firmware_id not in scan_results_store:
            raise HTTPException(status_code=404, detail="固件不存在")
        
        firmware_info = scan_results_store[firmware_id]
        firmware_path = firmware_info['path']
        
        # 使用队列进行扫描
        queue = get_queue()
        task_id = queue.add_task(firmware_path, firmware_type, firmware_id)
        
        # 等待完成（同步模式）
        task = queue.wait_for_completion(task_id, poll_interval=1.0)
        
        if not task:
            raise HTTPException(status_code=500, detail="任务创建失败")
        
        if task.status == TaskStatus.FAILED:
            raise HTTPException(status_code=500, detail=task.error_message)
        
        result = task.result
        
        scan_results_store[firmware_id]['result'] = result
        scan_results_store[firmware_id]['status'] = 'completed'
        
        return {"success": True, "result": result, "task_id": task_id}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"扫描失败：{e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/scan/batch")
async def scan_firmware_batch(request: Request):
    """批量扫描多个固件（接收 JSON）"""
    try:
        data = await request.json()
        files = data.get('files', [])
        
        if not files or len(files) == 0:
            raise HTTPException(status_code=400, detail="未提供文件列表")
        
        queue = get_queue()
        task_ids = []
        
        for file_info in files:
            path = file_info.get('path')
            firmware_type = file_info.get('type', 'bin')
            filename = file_info.get('filename', os.path.basename(path))
            
            if not path:
                continue
            
            task_id = queue.add_task(path, firmware_type, filename)
            task_ids.append({
                'task_id': task_id,
                'filename': filename,
                'status': 'queued'
            })
        
        logger.info(f"📦 批量提交 {len(task_ids)} 个扫描任务")
        
        return {
            "success": True,
            "submitted": len(task_ids),
            "tasks": task_ids
        }
        
    except Exception as e:
        logger.error(f"批量扫描提交失败：{e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
# 任务队列管理 API
# ============================================================

@app.get("/api/task/{task_id}")
async def get_task_status(task_id: str):
    """获取单个任务状态"""
    queue = get_queue()
    task = queue.get_task_status(task_id)
    
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    return TaskStatusResponse(**task.to_dict())


@app.get("/api/tasks")
async def list_tasks(
    status: Optional[str] = Query(default=None),
    limit: int = Query(default=100, le=500)
):
    """列出所有任务（可过滤状态）"""
    queue = get_queue()
    
    task_status = None
    if status:
        try:
            task_status = TaskStatus(status)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"无效的状态：{status}")
    
    tasks = queue.get_all_tasks(limit=limit, status=task_status)
    
    return {
        "total": len(tasks),
        "tasks": [task.to_dict() for task in tasks]
    }


@app.get("/api/queue/stats")
async def get_queue_stats():
    """获取队列统计信息"""
    queue = get_queue()
    stats = queue.get_queue_stats()
    return QueueStatsResponse(**stats)


@app.post("/api/task/{task_id}/cancel")
async def cancel_task(task_id: str):
    """取消任务"""
    queue = get_queue()
    success = queue.cancel_task(task_id)
    
    if not success:
        raise HTTPException(status_code=400, detail="无法取消该任务")
    
    return {"success": True, "message": "任务已取消"}


@app.delete("/api/tasks/clear")
async def clear_old_tasks(days: int = Query(default=7, ge=1)):
    """清理旧任务记录"""
    queue = get_queue()
    count = queue.clear_old_tasks(days)
    
    return {"success": True, "deleted_count": count, "days": days}


@app.get("/api/reports/{task_id}")
async def generate_yaml_report(task_id: str):
    """生成并下载 YAML 报告"""
    queue = get_queue()
    task = queue.get_task_status(task_id)
    
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    if task.status != TaskStatus.COMPLETED or not task.result:
        raise HTTPException(status_code=400, detail="任务未完成或无结果")
    
    result = task.result
    
    # 生成 YAML 报告
    report_data = f"""# 固件漏洞扫描报告
firmware:
  id: {result['firmware_id']}
  filename: {result['filename']}
  type: {result['firmware_type']}
  scan_time: {result['scan_time']}

summary:
  total_cves: {result['total_cves']}
  critical: {result['critical_count']}
  high: {result['high_count']}
  medium: {result['medium_count']}
  low: {result['low_count']}

r155_compliance:
  overall_score: {result['r155_compliance']['overall_score']:.2f}
  compliance_level: {result['r155_compliance']['compliance_level_text']}
  domains:
"""
    
    for domain, score in result['r155_compliance']['domain_scores'].items():
        report_data += f"    {domain}: {score:.1f}\n"
    
    report_data += "\ncomponents:\n"
    for comp in result['components']:
        report_data += f"- name: {comp['name']}\n"
        report_data += f"  version: {comp['version']}\n"
        report_data += f"  license: {comp['license']}\n\n"
    
    report_data += "\nvulnerabilities:\n"
    for vuln in result['vulnerabilities']:
        report_data += f"- cve_id: {vuln['cve_id']}\n"
        report_data += f"  component: {vuln['component']}\n"
        report_data += f"  severity: {vuln['severity']}\n"
        report_data += f"  cvss_score: {vuln['cvss_score']}\n"
        report_data += f"  priority_score: {vuln['priority_score']}\n"
        report_data += f"  r155_non_compliant: {vuln['r155_non_compliant']}\n\n"
    
    # 保存报告
    report_path = Path(config['paths']['reports']) / f"{task_id}.yaml"
    with open(report_path, 'w') as f:
        f.write(report_data)
    
    return FileResponse(report_path, media_type='text/yaml', filename=f"{task_id}_report.yaml")


@app.get("/api/compliance/{task_id}")
async def get_compliance_result(task_id: str):
    """获取 R155 合规检查结果"""
    queue = get_queue()
    task = queue.get_task_status(task_id)
    
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    if task.status != TaskStatus.COMPLETED or not task.result:
        raise HTTPException(status_code=400, detail="任务未完成或无结果")
    
    return task.result['r155_compliance']


@app.post("/api/report/r155-word")
async def generate_r155_word_report(task_id: str = Form(...)):
    """生成 R155 合规 Word 报告"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        queue = get_queue()
        task = queue.get_task_status(task_id)
        
        if not task or task.status != TaskStatus.COMPLETED:
            raise HTTPException(status_code=404, detail="任务不存在或未完成")
        
        result = task.result
        compliance = result['r155_compliance']
        
        # 创建文档
        doc = Document()
        
        # 标题
        title = doc.add_heading(f'EU R155 合规评估报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        doc.add_paragraph(f'\n固件名称：{compliance["firmware_name"]}')
        doc.add_paragraph(f'固件 ID: {compliance["firmware_id"]}')
        doc.add_paragraph(f'扫描时间：{compliance["scan_time"]}')
        
        # 综合评分
        doc.add_heading('\n📊 综合评分', level=1)
        p = doc.add_paragraph()
        p.add_run(f'总体得分：{compliance["overall_score"]:.2f}/100').bold = True
        
        # 合规等级
        level_color = {
            0: RGBColor(220, 53, 69),    # Red
            1: RGBColor(255, 193, 7),    # Yellow
            2: RGBColor(25, 135, 84),    # Green
            3: RGBColor(40, 167, 69),    # Dark Green
            4: RGBColor(0, 123, 255)     # Blue
        }
        
        level_p = doc.add_paragraph()
        level_run = level_p.add_run(f'合规等级：{compliance["compliance_level_text"]}')
        level_run.bold = True
        level_run.font.color.rgb = level_color.get(compliance["compliance_level"], RGBColor(0, 0, 0))
        
        # 域名得分表格
        doc.add_heading('\n🔐 领域得分', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '安全领域'
        hdr_cells[1].text = '得分'
        
        for domain, score in compliance['domain_scores'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = domain
            row_cells[1].text = f'{score:.1f}'
        
        # 高风险项目
        if compliance['high_risk_items']:
            doc.add_heading('\n⚠️ 高风险项目', level=1)
            for item in compliance['high_risk_items'][:10]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f'{item["type"]}: {item["id"]}').bold = True
                p.add_run(f' ({item["severity"]})')
                if len(item['description']) > 100:
                    p.add_run(f': {item["description"][:100]}...')
                else:
                    p.add_run(f': {item["description"]}')
        
        # 修复建议
        if compliance['remediation_recommendations']:
            doc.add_heading('\n✅ 修复建议', level=1)
            for rec in compliance['remediation_recommendations'][:10]:
                doc.add_paragraph(rec, style='List Number')
        
        # 统计数据
        doc.add_heading('\n📈 统计信息', level=1)
        stats = compliance['statistics']
        doc.add_paragraph(f'• 总 CVE 数：{stats["total_evidence"]}')
        doc.add_paragraph(f'• ✓ 合规项：{stats["compliant"]}')
        doc.add_paragraph(f'• ✗ 不合规项：{stats["non_compliant"]}')
        doc.add_paragraph(f'• ⚠ 部分合规：{stats["partial"]}')
        
        # 保存文件
        report_path = Path(config['paths']['reports']) / f"{task_id}_R155.docx"
        doc.save(report_path)
        
        return FileResponse(
            report_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"{task_id}_R155_Report.docx"
        )
        
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装，请运行：pip install python-docx")
    except Exception as e:
        logger.error(f"Word 报告生成失败：{e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/results/{firmware_id}")
async def get_results(firmware_id: str):
    """获取扫描结果"""
    if firmware_id not in scan_results_store:
        raise HTTPException(status_code=404, detail="扫描结果不存在")
    
    result = scan_results_store[firmware_id].get('result')
    if not result:
        raise HTTPException(status_code=404, detail="扫描尚未完成")
    
    return result


# @app.get("/api/dashboard") - 已删除重复定义
# async def dashboard(request):
#     ...


@app.post("/api/report/excel")
async def generate_excel_report(firmware_id: str = Form(...)):
    """生成 Excel 报告"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    
    if firmware_id not in scan_results_store:
        raise HTTPException(status_code=404, detail="扫描结果不存在")
    
    result = scan_results_store[firmware_id].get('result')
    if not result:
        raise HTTPException(status_code=404, detail="扫描尚未完成")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "漏洞报告"
    
    # 标题样式
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    # 表头
    headers = ["CVE ID", "组件", "版本", "严重程度", "CVSS", "优先级", "描述", "R155 合规"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    
    # 数据行
    severity_colors = {
        'critical': 'C72020',
        'high': 'FFC72C',
        'medium': 'FFFF00',
        'low': '92D050'
    }
    
    for row_idx, vuln in enumerate(result['vulnerabilities'], 2):
        color = severity_colors.get(vuln['severity'].lower(), 'FFFFFF')
        fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
        
        ws.cell(row=row_idx, column=1, value=vuln['cve_id'])
        ws.cell(row=row_idx, column=2, value=vuln['component'])
        ws.cell(row=row_idx, column=3, value=vuln['version'])
        
        sev_cell = ws.cell(row=row_idx, column=4, value=vuln['severity'])
        sev_cell.fill = fill
        
        ws.cell(row=row_idx, column=5, value=vuln['cvss_score'])
        ws.cell(row=row_idx, column=6, value=vuln['priority_score'])
        ws.cell(row=row_idx, column=7, value=vuln['description'])
        ws.cell(row=row_idx, column=8, value="❌ 不合规" if vuln['r155_non_compliant'] else "✅ 合规")
    
    # 调整列宽
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column].width = min(max_length + 2, 50)
    
    # 保存
    report_path = Path(config['paths']['reports']) / f"{firmware_id}_report.xlsx"
    wb.save(report_path)
    
    return FileResponse(report_path, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', filename=f"{firmware_id}_report.xlsx")


from report_generator.pdf_generator import generate_pdf_report


@app.post("/api/report/pdf")
async def generate_pdf_report(firmware_id: str = Form(...)):
    """生成 PDF 报告"""
    if firmware_id not in scan_results_store:
        raise HTTPException(status_code=404, detail="扫描结果不存在")
    
    result = scan_results_store[firmware_id].get('result')
    if not result:
        raise HTTPException(status_code=404, detail="扫描尚未完成")
    
    try:
        # 使用本地 PDF 生成器
        pdf_path = generate_pdf_report(firmware_id, result)
        return FileResponse(
            pdf_path, 
            media_type='application/pdf', 
            filename=f"{firmware_id}_security_report.pdf"
        )
    except Exception as e:
        logger.error(f"PDF 报告生成失败：{e}")
        raise HTTPException(status_code=500, detail=f"PDF 报告生成失败：{str(e)}")


@app.post("/api/report/word")
async def generate_word_report(firmware_id: str = Form(...)):
    """生成 Word 报告 (调用 Node.js 服务)"""
    import requests
    
    if firmware_id not in scan_results_store:
        raise HTTPException(status_code=404, detail="扫描结果不存在")
    
    result = scan_results_store[firmware_id].get('result')
    if not result:
        raise HTTPException(status_code=404, detail="扫描尚未完成")
    
    # 调用本地 Node.js 报告生成服务
    try:
        response = requests.post(
            "http://localhost:3000/api/report/word",
            json={'firmware_id': firmware_id, 'data': result},
            timeout=30
        )
        
        if response.status_code == 200:
            report_path = response.json()['path']
            return FileResponse(report_path, media_type='application/docx', filename=f"{firmware_id}_report.docx")
        else:
            raise HTTPException(status_code=500, detail="Word 报告生成失败")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/report/ppt")
async def generate_ppt_report(firmware_id: str = Form(...)):
    """生成 PPT 报告"""
    import requests
    
    if firmware_id not in scan_results_store:
        raise HTTPException(status_code=404, detail="扫描结果不存在")
    
    result = scan_results_store[firmware_id].get('result')
    if not result:
        raise HTTPException(status_code=404, detail="扫描尚未完成")
    
    try:
        response = requests.post(
            "http://localhost:3000/api/report/ppt",
            json={'firmware_id': firmware_id, 'data': result},
            timeout=30
        )
        
        if response.status_code == 200:
            report_path = response.json()['path']
            return FileResponse(report_path, media_type='application/vnd.ms-powerpoint', filename=f"{firmware_id}_report.pptx")
        else:
            raise HTTPException(status_code=500, detail="PPT 报告生成失败")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/scans")
async def list_scans():
    """列出所有扫描任务"""
    scans = []
    for fw_id, info in scan_results_store.items():
        scans.append({
            'id': fw_id,
            'filename': info['filename'],
            'upload_time': info['upload_time'],
            'status': info['status'],
            'total_cves': info.get('result', {}).get('total_cves', 0)
        })
    return JSONResponse(content=scans)


# ============================================================
# R155 合规报告 API
# ============================================================

@app.get("/api/compliance/{task_id}")
async def get_compliance_report(task_id: str):
    """获取 R155 合规报告
    
    Args:
        task_id: 扫描任务 ID
        
    Returns:
        R155 合规报告（包含评分、违规详情和建议）
    """
    queue = get_queue()
    task = queue.get_task_status(task_id)
    
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    if task.status != TaskStatus.COMPLETED or not task.result:
        raise HTTPException(
            status_code=400, 
            detail=f"任务状态：{task.status.value}，尚未完成或无结果"
        )
    
    compliance = task.result.get('r155_compliance', {})
    
    if not compliance:
        return {
            'error': '未找到合规报告',
            'message': '可能由于扫描结果较旧或不支持 R155 检查'
        }
    
    return compliance


@app.get("/api/compliance/categories/{task_id}")
async def get_compliance_by_category(task_id: str):
    """按类别查看 R155 合规得分"""
    queue = get_queue()
    task = queue.get_task_status(task_id)
    
    if not task or task.status != TaskStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="任务未完成")
    
    compliance = task.result.get('r155_compliance', {})
    category_scores = compliance.get('category_scores', {})
    
    # 返回格式化的类别得分
    result = []
    for category, score in category_scores.items():
        # 确定等级
        if score >= 85:
            level = "✅ 优秀"
        elif score >= 70:
            level = "⚠️ 良好"
        elif score >= 50:
            level = "⚠️ 中等"
        else:
            level = "❌ 需改进"
        
        result.append({
            'category': category,
            'score': score,
            'level': level
        })
    
    return sorted(result, key=lambda x: x['score'])


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host=config['server']['host'], port=config['server']['port'])
